import datetime
import os
from time import sleep

import docx
import requests
import wget
from FolderSearch import chazhao, chucun
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from win32com.client import gencache, constants

#########################################################################################
cid = {(76, 20),
       (17, 15),
       (95, 15),
       (75, 12),
       (182, 5),
       (124, 7),
       (86, 5),
       (85, 5),
       (21, 2),
       (71, 2),
       (122, 2),
       (138, 2),
       (161, 2),
       (172, 2),
       (183, 2),
       (184, 2)}


#########################################################################################

def huo_qu_shi_jian():
    p = int((datetime.datetime.now() - datetime.timedelta(days=7)).strftime("%Y%m%d"))
    n = int(datetime.datetime.now().strftime("%Y%m%d"))
    return p, n


def pa_qu_shi_ping_shu_ju(id, ye, p, n):
    url = 'https://s.search.bilibili.com/cate/search?main_ver=v3&search_type=video&view_type=hot_rank&order=click&copy_right=-1&cate_id=%d&page=%d&pagesize=100&jsonp=jsonp&time_from=%d&time_to=%d' % (
        id, ye, p, n)
    with requests.get(url) as t:
        c = t.json()
    return c


def shi_ping_shu_ju_zheng_li_he_fen_xi(c):
    if c['code'] != 0:
        return 'E:%d' % c['code']
    l = len(c['result'])
    nc = [{} for t in range(l)]
    for i in range(l):  # 'duration','id','author','play','pic','title','arcurl'
        if int(c['result'][i]['duration']) <= 600:
            nc[i]['score'] = int(pow(int(c['result'][i]['duration']), 1.5) * int(c['result'][i]['play']))
        else:
            nc[i]['score'] = int(
                15000 * int(c['result'][i]['play']) + pow(int(c['result'][i]['duration']) - 600, 0.635) * int(
                    c['result'][i]['play']))
        nc[i]['id'] = c['result'][i]['id']
        nc[i]['author'] = c['result'][i]['author']
        nc[i]['duration'] = c['result'][i]['duration']
        nc[i]['play'] = c['result'][i]['play']
        nc[i]['pic'] = c['result'][i]['pic']
        nc[i]['title'] = c['result'][i]['title']
        nc[i]['arcurl'] = c['result'][i]['arcurl']
        nc[i]['tag'] = c['result'][i]['tag']
    nnc = sorted(nc, key=lambda i: i['score'], reverse=True)
    return nnc


def shai_xuan(c):
    i = 0
    l = len(c)
    try:
        while i < l:
            if chazhao('DATA', str(c[i]['id'])):
                del c[i]
                i = i - 1
                l = l - 1
            i = i + 1
    except:
        pass
    return c


def shu_liang_kong_zhi(id, ye, c, n):
    l = len(c)
    if l < n:
        ye = ye + 1
        return yi_ci_all(id, ye, c, n)
    for i in range(n, l):
        del c[n]
    return c


def yi_ci_all(id, ye, cc, num):
    sleep(0.5)
    p, n = huo_qu_shi_jian()
    c = pa_qu_shi_ping_shu_ju(id, ye, p, n)
    con = shi_ping_shu_ju_zheng_li_he_fen_xi(c)
    try:
        con + ''
        print(id, ':', con)
        while 1:
            sleep(10000)
    except:
        pass
    ncon = shai_xuan(con)
    cc = cc + ncon
    cc = shu_liang_kong_zhi(id, ye, cc, num)
    chu_cun_id(cc)
    return cc


def chu_cun_id(cc):
    l = len(cc)
    for i in range(l):
        chucun('DATA', str(cc[i]['id']))


def chu_cun(c):
    doc = docx.Document()
    l = len(c)
    for i in range(l):
        url = c[i]['pic']
        str1 = '第%d个视频' % (i + 1)
        str2 = '链接：' + c[i]['arcurl'] + '\n' + '标题：' + c[i]['title'] + '\n' + '标签：' + c[i]['tag'] + '\n' + '作者：' + \
               c[i]['author'] + '\n' + '时长：' + \
               str(c[i]['duration']) + '(S)' + '\n' + '播放量：' + str(c[i]['play'])
        try:
            t = doc.add_paragraph()
            t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            t2 = doc.add_paragraph()
            t2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            t2.paragraph_format.left_indent = Inches(0.5)
            t3 = doc.add_paragraph()
            gs1 = t.add_run(str1)
            gs1.font.name = 'Times New Roman'
            gs1.font.element.rPr.rFonts.set(qn('w:eastAsia'), '新宋体')
            gs1.font.size = Pt(22)
            gs1.font.bold = True
            try:
                os.remove('T.jpg')
            except:
                pass
            wget.download('http://' + url, 'T.jpg')
            runp = p.add_run('')
            runp.add_picture('T.jpg', height=Inches(2.5))
            gs2 = t2.add_run(str2)
            gs2.font.name = 'Times New Roman'
            gs2.font.element.rPr.rFonts.set(qn('w:eastAsia'), '新宋体')
            gs2.font.size = Pt(12.5)
        except:
            pass
        if i + 1 != l:
            gs3 = t3.add_run('')
            gs3.add_break(docx.enum.text.WD_BREAK.PAGE)

    try:
        os.remove('T.jpg')
        os.remove(docxname)
    except:
        pass
    doc.save('TMP')


def pdf_hua():
    try:
        os.remove(pdfname)
    except:
        pass
    path = os.getcwd()
    wordpath = os.path.join(path, 'TMP')
    pdfpath = os.path.join(path, pdfname)
    word = gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(wordpath, ReadOnly=1)
    doc.ExportAsFixedFormat(pdfpath,
                            constants.wdExportFormatPDF,
                            Item=constants.wdExportDocumentWithMarkup,
                            CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
    word.Quit(constants.wdDoNotSaveChanges)


print('\n\n运行中...\n\n')


def run():
    c = []
    ccc = []
    for ID, N in cid:
        ccc = ccc + yi_ci_all(ID, 1, c, N)
    ccc = sorted(ccc, key=lambda i: i['score'], reverse=True)
    chu_cun(ccc)


docxname = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S") + '.docx'
pdfname = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S") + '.pdf'
try:
    run()
except:
    os.system('cls')
    print('\n\n程序运行出错')
    print('\n')
    os.system('pause')
try:
    pdf_hua()
    os.remove('TMP')
except:
    os.rename('TMP', docxname)
os.system('cls')
print('\n\n推荐结果已放在同目录的推荐文件中')
print('\n')
os.system('pause')
